import os
from docx import Document

import spacy
nlp = spacy.load("nl_core_news_lg")

ANON = '[ANONIEM]'
RESULTS_DIRECTORY="anonimized_files"


class MyDocuments:
    def __init__(self, load_dir: str):
        self.load_dir = load_dir
        self.document_paths = os.listdir(load_dir)

        self.documents = [Document(f"{load_dir}/{path}") for path in self.document_paths]

    def show_all_documents(self):
        print("[INIT] Found the following files: ")
        print("------------------------")
        for filename in self.document_paths:
            print(f"- {filename}")

        print("------------------------\n")

    def find_person_entities(self, full_text: str):
        """Use Spacy to detect names in string."""
        spacy_text = nlp(full_text)

        names = []
        for ent in spacy_text.ents:
            if ent.label_ == "PERSON" and ent.text not in names:
                names.append(ent.text)

        names.sort(key=len, reverse=True)
        print(f"[RUN] Found the following names: {names}\n")
        return names

    def anonimize_all_documents(self):
        """Go through all of the documents, and anonimize them."""
        if not os.path.exists(RESULTS_DIRECTORY):
            os.mkdir(RESULTS_DIRECTORY)

        for index, document in enumerate(self.documents):
            full_path = f"{RESULTS_DIRECTORY}/{self.document_paths[index]}"

            if not os.path.exists(full_path):
                full_text = self.get_full_text(document)
                names = self.find_person_entities(full_text)

                for name in names:
                    self.anonymize(document, name)

                self.documents[index].save(full_path)

            else:
                print(f"[ERROR] There is already a file named {self.document_paths[index]}, skipping.")


    def get_full_text(self, document: Document):
        full_text = ""
        for paragraph in document.paragraphs:
            full_text += f"{paragraph.text} "
        return full_text

    def replace_words(self, document: Document, word_to_find: str, replace_with: str):
        for paragraph in document.paragraphs:
            if word_to_find in paragraph.text:
                paragraph.text = paragraph.text.replace(word_to_find, replace_with)

    def anonymize(self, document: Document, word_to_remove: str):
        self.replace_words(document, word_to_remove, ANON)


